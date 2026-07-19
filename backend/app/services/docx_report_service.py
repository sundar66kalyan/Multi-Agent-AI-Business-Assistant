from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt


class DocxReportService:

    @staticmethod
    def generate(report_data: dict):

        finance = report_data["finance"]
        analytics = report_data["analytics"]

        reports = Path("reports")
        reports.mkdir(exist_ok=True)

        filename = (
            f"Business_Report_"
            f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )

        output = reports / filename

        doc = Document()

        title = doc.add_heading(
            "MULTI-AGENT AI BUSINESS ASSISTANT",
            level=1
        )

        title.runs[0].font.size = Pt(22)

        doc.add_heading("Finance Summary", level=2)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"

        rows = [
            ("Revenue", finance["revenue"]),
            ("Expenses", finance["expenses"]),
            ("Profit", finance["profit"]),
            ("Month", finance["month"]),
        ]

        for metric, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(metric)
            cells[1].text = str(value)

        doc.add_heading("System Analytics", level=2)

        table2 = doc.add_table(rows=1, cols=2)
        table2.style = "Table Grid"

        hdr2 = table2.rows[0].cells
        hdr2[0].text = "Metric"
        hdr2[1].text = "Value"

        rows = [
            ("Indexed Documents", analytics["documents"]),
            ("Chunks", analytics["chunks"]),
            ("Finance Records", analytics["finance_records"]),
        ]

        for metric, value in rows:
            cells = table2.add_row().cells
            cells[0].text = str(metric)
            cells[1].text = str(value)

        doc.save(output)

        return str(output)